import os
import logging
from typing import Dict, List, Optional, Union, Any
import tempfile
import shutil
from pathlib import Path
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)

class DocumentLoader:
    """
    A class for loading documents from various file formats.
    """
    
    def __init__(self):
        """Initialize the DocumentLoader instance."""
        logger.info("DocumentLoader initialized")
        self.supported_extensions = {
            # Text files
            '.txt': self._load_text,
            '.md': self._load_text,
            '.csv': self._load_text,
            '.json': self._load_text,
            '.xml': self._load_text,
            '.html': self._load_text,
            
            # Microsoft Office
            '.docx': self._load_docx,
            '.xlsx': self._load_text,
            '.pptx': self._load_text,
            
            # Programming languages
            '.py': self._load_code,
            '.cs': self._load_code,  # C#
            '.vb': self._load_code,  # Visual Basic
            '.fs': self._load_code,  # F#
            '.java': self._load_code,
            '.js': self._load_code,
            '.ts': self._load_code,
            '.cpp': self._load_code,
            '.c': self._load_code,
            '.h': self._load_code,
            '.hpp': self._load_code,
            
            # .NET specific
            '.csproj': self._load_code,
            '.sln': self._load_code,
            '.vbproj': self._load_code,
            '.fsproj': self._load_code,
            '.config': self._load_code,
            '.aspx': self._load_code,
            '.cshtml': self._load_code,
            '.xaml': self._load_code,
            '.dll': self._load_binary_stub,  # Just metadata for DLLs
            
            # Archives
            '.zip': self._load_archive,
            '.rar': self._load_archive,
            '.tar': self._load_archive,
            '.gz': self._load_archive,
            '.7z': self._load_archive,
            '.pdf': self._load_pdf,
        }
    
    def set_document_loader(self):
        """Set this loader as active."""
        return self
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document from the specified file path.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Dict containing document metadata and content.
        """
        try:
            logger.info(f"Loading document from {file_path}")
            
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return {"content": "[File not found]", "metadata": {"path": file_path, "size": 0, "type": "unknown"}}
            
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension in self.supported_extensions:
                loader_func = self.supported_extensions[file_extension]
                try:
                    content = loader_func(file_path)
                    if not content or content.strip() == "":
                        logger.warning(f"No content extracted from file: {file_path}")
                        content = "[No readable content found in the document]"
                except ImportError as e:
                    logger.error(f"Required library not installed for {file_extension}: {str(e)}")
                    content = f"[Error: Required library not installed for {file_extension} files]"
                except Exception as e:
                    logger.error(f"Error in loader function for {file_extension}: {str(e)}")
                    content = f"[Error loading {file_extension} file: {str(e)}]"
                
                metadata = {
                    "path": file_path,
                    "size": os.path.getsize(file_path),
                    "type": file_extension[1:],  # Remove the dot
                    "error": None if content and not content.startswith("[Error") else content
                }
                return {
                    "content": content,
                    "metadata": metadata,
                    "extension": file_extension  # Add extension to the returned data
                }
            else:
                logger.warning(f"Unsupported file extension: {file_extension}")
                # Try to load as text anyway as a fallback
                try:
                    content = self._load_text(file_path)
                    metadata = {
                        "path": file_path,
                        "size": os.path.getsize(file_path),
                        "type": "unknown",
                        "error": None if content and not content.startswith("[Error") else content
                    }
                    return {"content": content, "metadata": metadata}
                except Exception as inner_e:
                    logger.error(f"Failed to load file as text: {str(inner_e)}")
                    return {
                        "content": "[Unsupported file format]",
                        "metadata": {
                            "path": file_path,
                            "size": os.path.getsize(file_path),
                            "type": "unknown",
                            "error": f"Failed to load file as text: {str(inner_e)}"
                        }
                    }
        
        except Exception as e:
            logger.error(f"Error loading document: {str(e)}")
            return {
                "content": f"[Error loading document: {str(e)}]",
                "metadata": {
                    "path": file_path,
                    "size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
                    "type": "unknown",
                    "error": str(e)
                }
            }
    
    def _determine_file_type(self, extension: str) -> str:
        """Determine the general file type based on extension."""
        code_extensions = ['.py', '.cs', '.vb', '.fs', '.java', '.js', '.ts', '.cpp', '.c', '.h', '.hpp',
                          '.csproj', '.sln', '.vbproj', '.fsproj', '.config', '.aspx', '.cshtml', '.xaml']
        
        if extension in ['.txt', '.md']:
            return "text"
        elif extension in ['.csv']:
            return "spreadsheet"
        elif extension in ['.docx']:
            return "document"
        elif extension in ['.xlsx']:
            return "spreadsheet"
        elif extension in ['.pptx']:
            return "presentation"
        elif extension in code_extensions:
            return "code"
        elif extension in ['.zip', '.rar', '.tar', '.gz', '.7z']:
            return "archive"
        else:
            return "unknown"
    
    def _load_text(self, file_path: str) -> str:
        """Load content from a text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with a different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading text file with latin-1 encoding: {str(e)}")
                # Try binary mode as a last resort
                with open(file_path, 'rb') as file:
                    binary_content = file.read()
                    try:
                        return binary_content.decode('utf-8', errors='replace')
                    except:
                        return f"[Binary file content - {len(binary_content)} bytes]"
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            return f"[Error reading text file: {str(e)}]"
    
    def _load_docx(self, file_path: str) -> str:
        """Load content from a .docx file."""
        try:
            from docx import Document
            doc = Document(file_path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    content.append(para.text)
            
            # Also try to get text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)
            
            text_content = "\n".join(content)
            if not text_content.strip():
                return "[No readable text found in the document]"
            return text_content
        except ImportError:
            logger.error("python-docx not installed")
            return "[Error: python-docx library required to process .docx files]"
        except Exception as e:
            logger.error(f"Error reading .docx file: {str(e)}")
            return f"[Error reading .docx file: {str(e)}]"
    
    def _load_code(self, file_path: str) -> str:
        """Load content from a code file."""
        return self._load_text(file_path)
    
    def _load_binary_stub(self, file_path: str) -> str:
        """Create a stub description for binary files like DLLs."""
        file_size = os.path.getsize(file_path)
        file_name = os.path.basename(file_path)
        
        return f"[Binary file: {file_name}, Size: {file_size} bytes]"
    
    def _load_archive(self, file_path: str) -> str:
        """
        Extract and process archive files.
        For archives, we create a summary of their contents.
        """
        try:
            import zipfile
            import tarfile
            import rarfile  # You might need to install this package
            
            file_extension = os.path.splitext(file_path)[1].lower()
            file_name = os.path.basename(file_path)
            result = f"Archive: {file_name}\n\nContents:\n"
            
            # Create a temporary directory for extraction
            temp_dir = tempfile.mkdtemp()
            
            try:
                # Extract based on file type
                if file_extension == '.zip':
                    with zipfile.ZipFile(file_path, 'r') as zip_ref:
                        file_list = zip_ref.namelist()
                        # Only extract if it's a small archive
                        if len(file_list) < 100:  # Arbitrary limit to prevent huge extractions
                            zip_ref.extractall(temp_dir)
                
                elif file_extension in ['.tar', '.gz']:
                    with tarfile.open(file_path, 'r:*') as tar_ref:
                        file_list = tar_ref.getnames()
                        # Only extract if it's a small archive
                        if len(file_list) < 100:
                            tar_ref.extractall(temp_dir)
                
                elif file_extension == '.rar':
                    with rarfile.RarFile(file_path) as rar_ref:
                        file_list = rar_ref.namelist()
                        # Only extract if it's a small archive
                        if len(file_list) < 100:
                            rar_ref.extractall(temp_dir)
                
                else:
                    return f"[Archive file: {file_name}, unsupported extraction]"
                
                # Add file list to result
                for i, file in enumerate(sorted(file_list)[:100]):  # Limit display to 100 files
                    result += f"- {file}\n"
                
                if len(file_list) > 100:
                    result += f"... and {len(file_list) - 100} more files\n"
                
                # Process common file types within the archive
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        if file.endswith(('.py', '.cs', '.md', '.txt', '.java')):
                            full_path = os.path.join(root, file)
                            rel_path = os.path.relpath(full_path, temp_dir)
                            
                            # Only include relatively small files
                            if os.path.getsize(full_path) < 100000:  # 100KB limit
                                try:
                                    with open(full_path, 'r', encoding='utf-8') as f:
                                        content = f.read()
                                    result += f"\n\n--- File: {rel_path} ---\n{content}\n"
                                except Exception as e:
                                    result += f"\n\n--- File: {rel_path} ---\n[Error reading file: {str(e)}]\n"
            
            finally:
                # Clean up the temporary directory
                shutil.rmtree(temp_dir)
            
            return result
            
        except ImportError as e:
            return f"[Archive file: {os.path.basename(file_path)}, extraction unavailable: {str(e)}]"
        except Exception as e:
            logger.error(f"Error processing archive {file_path}: {str(e)}")
            return f"[Archive file: {os.path.basename(file_path)}, error during extraction: {str(e)}]"

    def _load_pdf(self, file_path: str) -> str:
        """Load content from a PDF file."""
        try:
            reader = PdfReader(file_path)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return "\n".join(text).strip() or "[No text found in PDF]"
        except ImportError:
            logger.error("PyPDF2 not installed")
            return "[Error: PyPDF2 library required to process PDF files]"
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            return f"[Error reading PDF file: {str(e)}]"